import os
from datetime import date, timedelta

from aiogram import types, Dispatcher
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from sqlalchemy import select
import traceback
from database.db import async_session_maker
from database.models import User, Deal, Task, TaskStatus, Report
import torch
from handlers.reports.ai_model import tokenizer, model
from handlers.reports.generators import (
    admin_performance,
    admin_deals,
    admin_sales,
    admin_funnel,
    admin_timeline,
    admin_tables
)

# -------------------------------------------------
# 🤖 ИИ-РЕКОМЕНДАЦИИ
# -------------------------------------------------
def generate_ai_recommendation(stats: dict) -> str:
    """
    Генерация списка рекомендаций по работе менеджера.
    На выходе только готовый текст рекомендаций, без промта.
    """
    prompt = (
        f"Ты — опытный бизнес-аналитик CRM-системы. "
        f"Используй только фактические данные менеджера: "
        f"Новые: {stats['new']}, В работе: {stats['in_progress']}, "
        f"Приостановленные: {stats['on_hold']}, Закрытые: {stats['completed']}, Всего: {stats['total']}. "
        f"Составь 3–4 конкретные деловые рекомендации для менеджера, чтобы повысить эффективность работы. "
        f"Выдавай только связный текст абзацем, без заголовков и промтов."
    )

    try:
        inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=512)
        with torch.no_grad():
            output = model.generate(
                **inputs,
                max_new_tokens=150,
                do_sample=True,
                temperature=0.3,
                top_p=0.85,
                repetition_penalty=1.1,
                pad_token_id=tokenizer.eos_token_id
            )

        # Декодируем и чистим вывод
        text = tokenizer.decode(output[0], skip_special_tokens=True).strip()

        # Убираем случайно вставленные части промта (если они есть)
        if prompt in text:
            text = text.replace(prompt, "").strip()

        # Если текст пустой, выдаём заглушку
        if not text:
            return (
                "Рекомендуется контролировать новые и приостановленные сделки, "
                "оптимизировать распределение задач и завершать приоритетные сделки."
            )

        # Возвращаем уже готовый текст рекомендаций
        return text

    except Exception as e:
        print(f"⚠️ Ошибка ИИ: {e}")
        return (
            "Рекомендуется контролировать новые и приостановленные сделки, "
            "оптимизировать распределение задач и завершать приоритетные сделки."
        )

# -------------------------------------------------
# 📈 ПРОГНОЗ НА СЛЕДУЮЩИЙ ГОД
# -------------------------------------------------
def generate_forecast(stats: dict) -> str:
    """
    Программный прогноз на следующий год, базируется на текущих данных.
    """
    growth_factor = 1.1  # прогнозируемый рост 10%
    forecast_new = int(stats['new'] * growth_factor)
    forecast_in_progress = int(stats['in_progress'] * growth_factor)
    forecast_on_hold = int(stats['on_hold'] * growth_factor)
    forecast_completed = int(stats['completed'] * growth_factor)
    forecast_total = forecast_new + forecast_in_progress + forecast_on_hold + forecast_completed

    return (
        f"На следующий год ожидается рост новых сделок до {forecast_new}, "
        f"текущих сделок в работе до {forecast_in_progress}, "
        f"приостановленных сделок до {forecast_on_hold}, "
        f"закрытых сделок до {forecast_completed}, "
        f"всего задач до {forecast_total}. "
        "Рекомендуется контролировать распределение ресурсов и приоритетов для эффективного выполнения задач."
    )

# -------------------------------------------------
# 📊 ВЫВОД ПО ВСЕМ ДАННЫМ (без ИИ)
# -------------------------------------------------
def generate_conclusion(stats: dict, deal_stats: dict, task_stats: dict) -> str:
    """
    Генерация связного заключения по всем данным отчета:
    сделки, задачи, диаграммы.
    """
    conclusion = (
        f"В течение периода менеджеры успешно обработали основные сделки и задачи. "
        f"Общее количество сделок составило {stats['total']}, из которых {deal_stats['Новая']} новых, "
        f"{deal_stats['В работе']} в работе, {deal_stats['Приостановлена']} приостановленных и "
        f"{deal_stats['Закрыта']} закрытых. "
        f"Команда выполнила {task_stats['total_created']} задач, из которых {task_stats['total_done']} завершено, "
        f"{task_stats['total_overdue']} просрочено. "
        f"Диаграммы и таблицы отчета показывают распределение нагрузки и эффективность работы, "
        f"позволяя выявить узкие места и оптимизировать процесс управления проектами. "
        f"Общий уровень эффективности можно считать стабильным с потенциалом роста при дальнейшем контроле приоритетов."
    )
    return conclusion

# -------------------------------------------------
# 🖼 ЗАМЕНА ДИАГРАММ (ОТДЕЛЬНО!)
# -------------------------------------------------
def replace_diagram_placeholders(doc, diagram_map: dict):
    print("🖼 Начинаю замену диаграмм/таблиц")
    for p_idx, paragraph in enumerate(doc.paragraphs):
        for placeholder, image_path in diagram_map.items():
            if placeholder in paragraph.text:
                print(f"🖼 Найден плейсхолдер {placeholder} в абзаце {p_idx}")
                paragraph.text = ""
                run = paragraph.add_run()
                if os.path.exists(image_path):
                    run.add_picture(image_path, width=Inches(6))
                    print(f"✅ Вставлено изображение: {image_path}")
                else:
                    run.add_text(f"[Изображение {placeholder} недоступно]")
                    print(f"⚠️ Файл не найден: {image_path}")


# -------------------------------------------------
# 📝 ЗАМЕНА ТЕКСТА С СОХРАНЕНИЕМ СТИЛЯ
# -------------------------------------------------
def replace_text_placeholders_preserve_style(doc, replacements: dict):
    print("📝 Начинаю замену текстовых плейсхолдеров")
    for p_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.runs:
            continue
        original_text = paragraph.text
        new_text = original_text
        for key, value in replacements.items():
            if key in new_text:
                print(f"📝 Абзац {p_idx}: {key} → {value}")
                new_text = new_text.replace(key, str(value))
        if new_text != original_text:
            first_run = paragraph.runs[0]
            style = {
                "bold": first_run.bold,
                "italic": first_run.italic,
                "underline": first_run.underline,
                "font_name": first_run.font.name,
                "font_size": first_run.font.size,
                "font_color": first_run.font.color.rgb,
            }
            for run in paragraph.runs:
                run.text = ""
            run = paragraph.add_run(new_text)
            run.bold = style["bold"]
            run.italic = style["italic"]
            run.underline = style["underline"]
            if style["font_name"]:
                run.font.name = style["font_name"]
            if style["font_size"]:
                run.font.size = style["font_size"]
            if style["font_color"]:
                run.font.color.rgb = style["font_color"]
            print(f"✅ Абзац {p_idx} обновлён со стилем")


# -------------------------------------------------
# 📌 КНОПКА «ОТЧЁТ»
# -------------------------------------------------
async def generate_report_cb_handler(query: types.CallbackQuery):
    await query.answer("📊 Формирование отчёта")
    from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
    kb = InlineKeyboardMarkup(row_width=2)
    kb.add(
        InlineKeyboardButton("За неделю", callback_data="report_period_week"),
        InlineKeyboardButton("За месяц", callback_data="report_period_month"),
        InlineKeyboardButton("За год", callback_data="report_period_year"),
    )
    await query.message.answer("Выберите период отчёта:", reply_markup=kb)



# -------------------------------------------------
# 📌 ГЕНЕРАЦИЯ ОТЧЁТА
# -------------------------------------------------
async def report_period_cb_handler(query: types.CallbackQuery):
    try:
        period_days = {
            "report_period_week": 7,
            "report_period_month": 30,
            "report_period_year": 365,
        }.get(query.data, 30)

        end_date = date.today()
        start_date = end_date - timedelta(days=period_days)
        print(f"📅 Период отчёта: {start_date} → {end_date}")

        async with async_session_maker() as session:
            admin = await session.scalar(select(User).where(User.telegram_id == query.from_user.id))
            admin_name = admin.full_name if admin else "Администратор"
            print(f"👤 Администратор отчёта: {admin_name}")

            # -------------------------
            # СДЕЛКИ И ЗАДАЧИ
            # -------------------------
            deals = (await session.execute(select(Deal))).scalars().all()
            tasks = (await session.execute(select(Task))).scalars().all()

            deal_stats = {"Новая": 0, "В работе": 0, "Приостановлена": 0, "Закрыта": 0}
            for d in deals:
                if d.stage:
                    deal_stats[d.stage.value] += 1
            total_deals = sum(deal_stats.values())

            total_created = len(tasks)
            total_done = sum(1 for t in tasks if t.status == TaskStatus.done)
            total_overdue = sum(1 for t in tasks if t.status == TaskStatus.overdue)
            employees = {t.id_employee for t in tasks if t.id_employee}
            avg_load = round(total_created / max(len(employees), 1), 1)

            stats = {
                "new": deal_stats["Новая"],
                "in_progress": deal_stats["В работе"],
                "on_hold": deal_stats["Приостановлена"],
                "completed": deal_stats["Закрыта"],
                "total": total_deals
            }

            task_stats = {
                "total_created": total_created,
                "total_done": total_done,
                "total_overdue": total_overdue,
                "avg_load": avg_load
            }

        label = f"{period_days}d"

        # -------------------------
        # Диаграммы и таблицы
        # -------------------------
        await admin_performance.generate_admin_performance_diagram(start_date, end_date, label)
        await admin_deals.generate_admin_deals_diagram(start_date, end_date, label)
        await admin_sales.generate_admin_sales_diagram(start_date, end_date, label)
        await admin_funnel.generate_admin_sales_funnel(start_date, end_date, label)
        await admin_timeline.generate_admin_tasks_timeline_diagram(start_date, end_date, label)
        sales_table_path = await admin_tables.generate_admin_sales_table(start_date, end_date, label)
        performance_table_path = await admin_tables.generate_admin_performance_table(start_date, end_date, label)

        # -------------------------
        # WORD → PDF
        # -------------------------
        template = "reports/admin_report_template.docx"
        output_dir = "reports/generated"
        os.makedirs(output_dir, exist_ok=True)
        docx_path = f"{output_dir}/admin_report_{label}.docx"
        pdf_path = f"{output_dir}/admin_report_{label}.pdf"

        doc = Document(template)

        # диаграммы
        replace_diagram_placeholders(doc, {
            "{{diagram_admin_performance}}": f"reports/images/admin_performance_report_{label}.png",
            "{{diagram_admin_deals}}": f"reports/images/admin_deals_progress_{label}.png",
            "{{diagram_admin_sales}}": f"reports/images/admin_sales_by_clients_{label}.png",
            "{{diagram_admin_funnel}}": f"reports/images/admin_sales_funnel_{label}.png",
            "{{diagram_admin_timeline}}": f"reports/images/admin_tasks_timeline_{label}.png",
            "{{table_admin_sales}}": sales_table_path,
            "{{table_admin_performance}}": performance_table_path
        })

        # текст
        replace_text_placeholders_preserve_style(doc, {
            "{admin_name}": admin_name,
            "{date_start}": start_date.strftime("%d.%m.%Y"),
            "{date_end}": end_date.strftime("%d.%m.%Y"),
            "{total_deals}": total_deals,
            "{new_deals}": deal_stats["Новая"],
            "{in_progress_deals}": deal_stats["В работе"],
            "{on_hold_deals}": deal_stats["Приостановлена"],
            "{completed_deals}": deal_stats["Закрыта"],
            "{conversion_rate}": round(deal_stats["Закрыта"] / total_deals * 100, 1) if total_deals else 0,
            "{total_created}": total_created,
            "{total_done}": total_done,
            "{total_overdue}": total_overdue,
            "{avg_load}": avg_load,
            "{recommendation1}": generate_ai_recommendation(stats),
            "{recommendation2}": generate_ai_recommendation(stats),
            "{recommendation3}": generate_ai_recommendation(stats),
            "{data_forecast}": generate_forecast(stats),
            "{conclusion}": generate_conclusion(stats, deal_stats, task_stats),
        })

        doc.save(docx_path)
        convert(docx_path, pdf_path)

        await query.message.answer_document(
            types.InputFile(pdf_path),
            caption=f"📄 Отчёт за {period_days} дней сформирован",
            parse_mode=None
        )

        print("✅ Отчёт успешно отправлен")

    except Exception as e:
        print("⚠️ Ошибка генерации отчёта:")
        traceback.print_exc()
        await query.message.answer(
            "⚠️ Произошла ошибка при формировании отчёта.\nПодробности смотри в логах."
        )

    ai_summary_text = generate_ai_recommendation(stats)

    async with async_session_maker() as session:
        report = Report(
            report_name=f"Отчёт администратора за {period_days} дней",
            report_type='ai_analysis',  # или 'summary', можно динамически менять
            generated_by=admin.id_user if admin else None,
            ai_summary=ai_summary_text
        )
        session.add(report)
        await session.commit()
        print(f"💾 Отчёт сохранён в БД с id {report.id_report}")


def register_admin_generate_report(dp: Dispatcher):
    dp.register_callback_query_handler(generate_report_cb_handler, lambda c: c.data == "report")
    dp.register_callback_query_handler(report_period_cb_handler, lambda c: c.data.startswith("report_period_"))
    print("✅ admin_generate_report зарегистрирован")
